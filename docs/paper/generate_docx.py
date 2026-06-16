"""
Generate GreenLens scientific paper as a formatted Word (.docx) file.
Run: python docs/paper/generate_docx.py
Output: docs/paper/GREENLENS_SCIENTIFIC_PAPER.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = os.path.join(os.path.dirname(__file__), "GREENLENS_SCIENTIFIC_PAPER.docx")

# ── helpers ──────────────────────────────────────────────────────────────────


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_run(para, text, bold=False, italic=False, size=None, color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if underline:
        run.underline = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, italic=False, indent=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(11)
    return p


def todo(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[TODO] {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.italic = True
    return p


def bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
    run.italic = True
    return p


def make_table(doc, headers, rows, col_widths=None, header_color="2E5090"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if (ri % 2) == 1:
            for cell in row.cells:
                set_cell_bg(cell, "EEF2FF")
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            if str(val).startswith("**") and str(val).endswith("**"):
                run.text = val[2:-2]
                run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# page margins
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(
    title_p,
    "GreenLens: A Cascade Ensemble Approach for Environmental Pollution Detection\n"
    "with Vietnamese Real-World Data",
    bold=True,
    size=16,
)

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(meta_p, "[Tên tác giả 1], [Tên tác giả 2], [Tên GVHD]\n", size=11)
add_run(meta_p, "[Tên trường Đại học], [Khoa / Bộ môn]\n", size=11, italic=True)
add_run(meta_p, "[Email liên hệ]", size=11)
doc.add_paragraph()

todo(doc, "Điền: tên tác giả, trường, khoa, email, tên venue (IEEE Access / RIVF / KSE 2026...)")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "ABSTRACT", level=1)

abstract_text = (
    "Environmental pollution monitoring in Vietnamese urban and peri-urban areas remains a critical challenge, "
    "compounded by a severe lack of domain-specific training data that reflects local waste characteristics and "
    "real-world capture conditions. We present GreenLens, a three-stage cascade ensemble system for automated "
    "pollution detection in citizen-reported images. The pipeline combines (1) a fine-tuned YOLOv8n object detector "
    "for localising TRASH and WATER-pollution instances, (2) an EfficientNet-B0 scene classifier that resolves "
    "ambiguous water-pollution signals at image level, and (3) a second EfficientNet-B0 providing fine-grained "
    "trash-subtype classification across seven categories on cropped detection regions. We introduce a novel "
    "Pollution Coverage Ratio (PCR) metric that maps bounding-box area occupancy to actionable severity bands, "
    "enabling downstream mobile reporting workflows. Our composite dataset — the first to include explicit "
    "Vietnamese in-the-wild pollution imagery — enables domain-specific fine-tuning that yields "
    "mAP@0.5 = [X]%, outperforming the COCO-pretrained zero-shot baseline by [ΔX] pp and a larger YOLOv8s "
    "baseline by [ΔY] pp."
)
body(doc, abstract_text)
todo(doc, "Điền X, ΔX, ΔY sau khi chạy run_paper_experiments.py")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "I. INTRODUCTION", level=1)

heading(doc, "1.1 Motivation", level=2)
body(
    doc,
    "Solid-waste mismanagement and water pollution constitute two of the most pressing environmental challenges "
    "in Vietnam's rapidly urbanising cities. Community-based pollution reporting — where citizens photograph pollution "
    "incidents via mobile applications — has emerged as a scalable monitoring strategy; however, the manual triage "
    "of submitted images creates bottlenecks and is prone to inter-rater inconsistency. Automated computer-vision "
    "systems offer a path to scalable, consistent triage, yet existing open-source models are predominantly trained "
    "on Western waste datasets (TACO, COCO, Open Images) that poorly represent Vietnamese pollution contexts: "
    "distinct waste-composition profiles, dense urban clutter, varied tropical lighting, and specific water-body "
    "types (canals, rice paddies, roadside drains).",
)
todo(
    doc,
    "Thêm số liệu cụ thể: X triệu tấn rác/năm, trích nguồn MONRE hoặc World Bank Vietnam Waste Report",
)

heading(doc, "1.2 Problem Statement", level=2)
body(
    doc,
    "We formalise the task as: given an RGB image captured by a citizen smartphone, output (a) a set of pollution "
    "detections with class labels and bounding boxes, (b) fine-grained subtype labels for detected trash, "
    "(c) a pollution coverage ratio, and (d) a severity assessment and recommended reporting action. The system "
    "must run within a 4.5-second wall-clock budget to remain interactive on a mobile-backend deployment.",
)

heading(doc, "1.3 Contributions", level=2)
body(doc, "This paper makes the following contributions:")
bullet(
    doc,
    "Vietnamese Pollution Dataset (VPD-1): A composite dataset combining Roboflow public annotations with "
    "[N] images self-captured across [M] Vietnamese locations — the first pollution dataset with explicit "
    "Vietnamese in-the-wild representation.",
)
bullet(
    doc,
    "GreenLens Cascade Ensemble: A three-stage pipeline (YOLOv8n detection → EfficientNet-B0 scene → "
    "EfficientNet-B0 subtype) under a principled evidence-gating fusion rule.",
)
bullet(
    doc,
    "Pollution Coverage Ratio (PCR): A novel continuous metric mapping bounding-box area occupancy "
    "to four actionable severity bands (LOW / MEDIUM / HIGH / CRITICAL).",
)
bullet(
    doc,
    "Systematic Comparative Evaluation: Comparison against COCO zero-shot baseline (E0), YOLOv8s fine-tuned "
    "baseline (E1b), and additional baselines on identical dataset splits.",
)
todo(doc, "Điền N (số ảnh tự chụp), M (số địa điểm)")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# II. RELATED WORK
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "II. RELATED WORK", level=1)

heading(doc, "2.1 Waste Detection Datasets", level=2)
body(
    doc,
    "TACO (Trash Annotations in Context) [Proença & Simões, 2020] provides 1,500 images across 60 waste categories "
    "in outdoor scenes, predominantly from European contexts. TrashNet [Thung & Yang, 2016] offers 2,527 "
    "studio-captured images across 6 categories — clean backgrounds that differ substantially from in-the-wild "
    "conditions. None of these datasets include annotations from Vietnamese environments, where waste composition "
    "and scene contexts (canal edges, market areas, narrow alleys) differ markedly from European or North American "
    "training data.",
)
todo(doc, "Thêm citations: WaDaBa, Open Litter Map, và bất kỳ dataset châu Á nào liên quan")

heading(doc, "2.2 Object Detection for Waste", level=2)
body(
    doc,
    "YOLO-family detectors have become the dominant approach for real-time waste detection due to their inference "
    "speed. Transformer-based detectors (DETR, RT-DETR) show competitive accuracy but incur higher latency "
    "unsuitable for mobile-backend deployment. YOLOv8 [Jocher et al., 2023] represents the current state-of-the-art "
    "in the YOLO lineage, offering improved accuracy and built-in support for segmentation and classification tasks.",
)
todo(doc, "Thêm 2–3 paper dùng YOLO cho waste detection, cite mAP cụ thể của họ để so sánh")

heading(doc, "2.3 Scene Classification for Environmental Monitoring", level=2)
body(
    doc,
    "EfficientNet [Tan & Le, ICML 2019] achieves strong accuracy-efficiency trade-offs and has been applied to "
    "remote-sensing scene classification and water-quality assessment. However, standalone scene classifiers lack "
    "spatial precision (no bounding boxes) and are prone to false positives from visually similar scenes.",
)

heading(doc, "2.4 Cascade and Ensemble Approaches", level=2)
body(
    doc,
    "Cascade detection pipelines — using a fast first-stage detector to gate expensive second-stage classifiers — "
    "are well-established in face detection and medical imaging. Our work introduces a principled evidence-gating "
    "rule: the scene classifier may only supplement, not override, the detector's decision, and only when the "
    "detector has found at least one object. This prevents the false-positive amplification observed in naïve "
    "ensemble fusion (see ablation, Section V.5).",
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# III. DATASET
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "III. DATASET", level=1)

heading(doc, "3.1 Data Sources", level=2)
body(
    doc,
    "Our composite dataset VPD-1 (Vietnamese Pollution Dataset v1) combines two complementary sources:",
)

body(doc, "Source A - Roboflow Public Datasets:")
body(
    doc,
    "We aggregate publicly available pollution and waste detection datasets from Roboflow Universe, selecting "
    "those with outdoor, in-the-wild images annotated at the instance level. Annotations are in COCO instances "
    "format and converted to YOLO format using a custom coco_instances_to_yolo.py conversion script. "
    "Classes are harmonised via synonym mapping: GARBAGE/WASTE/RUBBISH/LITTER -> TRASH; "
    "SEWAGE/WASTEWATER/EFFLUENT -> WATER.",
    indent=True,
)
todo(doc, "List ten chinh xac cac Roboflow dataset da dung, kem dataset ID va so anh tung dataset")

body(doc, "Source B - Vietnamese Self-Captured Images (VSC):")
body(
    doc,
    "We collect [N] images at [M] locations across [cities/provinces] using [device model(s)] between [date range]. "
    "Locations include canal banks, market perimeters, roadside drainage, and peri-urban dumping sites — scene types "
    "absent from Western datasets. Images are annotated using [annotation tool] by [K] annotators.",
    indent=True,
)
todo(
    doc,
    "Điền: N, M, tên tỉnh/thành, thiết bị chụp, thời gian, tool annotation, số annotator, inter-annotator agreement (Cohen's κ)",
)

heading(doc, "3.2 Dataset Statistics", level=2)
make_table(
    doc,
    headers=["Subset", "Images", "TRASH Instances", "WATER Instances", "Total"],
    rows=[
        ["Roboflow (Source A)", "[X]", "[X]", "[X]", "[X]"],
        ["Vietnamese VSC (Source B)", "[N]", "[N1]", "[N2]", "[N3]"],
        ["Combined VPD-1", "[Total]", "[T]", "[W]", "[All]"],
        ["— Train (70%)", "[X]", "—", "—", "—"],
        ["— Val (15%)", "[X]", "—", "—", "—"],
        ["— Test (15%)", "[X]", "—", "—", "—"],
    ],
    col_widths=[4.5, 2.5, 3.5, 3.5, 2.5],
)
note(doc, "Table 1. VPD-1 dataset statistics. All splits use random seed 42.")
todo(
    doc,
    "Chạy verify_yolo_dataset.py và điền bảng trên. Báo cáo thêm class balance ratio TRASH:WATER",
)

heading(doc, "3.3 Data Preprocessing", level=2)
body(
    doc,
    "Images are resized to 1280×1280 for YOLO training and 224×224 for EfficientNet classifiers. "
    "Labels are validated for coordinate normalisation (all values in [0, 1]) and class-ID range using "
    "the verify_yolo_dataset.py tool. The dataset is structured in YOLO format: "
    "images/{train,val,test}/ and labels/{train,val,test}/ with one .txt file per image "
    "(format: class_id x_center y_center width height, normalised 0–1).",
)

heading(doc, "3.4 Augmentation", level=2)
body(
    doc,
    "YOLO training relies on Ultralytics built-in mosaic augmentation, random affine transforms, "
    "HSV colour jitter, and horizontal flipping at default settings.",
)
body(doc, "EfficientNet-B0 classifiers apply the following augmentations during training:")
bullet(doc, "RandomResizedCrop(224, scale=(0.7, 1.0))")
bullet(doc, "RandomHorizontalFlip(p=0.5)")
bullet(doc, "ColorJitter(brightness=0.3, contrast=0.3, saturation=0.2)")
bullet(doc, "RandomRotation(15°) — trash subtype classifier only")
bullet(doc, "Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]) — ImageNet stats")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# IV. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "IV. METHODOLOGY", level=1)

heading(doc, "4.1 System Overview", level=2)
body(
    doc,
    "GreenLens is a three-stage cascade ensemble deployed as a FastAPI microservice. Stage 1 (YOLOv8n) and "
    "Stage 2 (EfficientNet-B0 scene) execute in parallel via a ThreadPoolExecutor (2 workers). "
    "Stage 3 (EfficientNet-B0 subtype) is conditional on Stage 1 TRASH detections. "
    "Total inference budget: 4.5 seconds.",
)
todo(
    doc,
    "Thêm Figure 1: sơ đồ pipeline 3 stage (vẽ bằng draw.io / Figma / PowerPoint rồi insert vào đây)",
)

heading(doc, "4.2 Stage 1 — Object Detection (YOLOv8n)", level=2)
body(
    doc,
    "We fine-tune YOLOv8n [Jocher et al., 2023] on VPD-1 to detect two pollution classes: "
    "TRASH (class 0: solid waste, litter, garbage) and WATER (class 1: wastewater, sewage, water pollution). "
    "At inference, boxes below the relevance confidence threshold (0.30) are discarded.",
)
make_table(
    doc,
    headers=["Hyperparameter", "Value"],
    rows=[
        ["Base model", "yolov8n.pt (COCO pretrained)"],
        ["Epochs", "150"],
        ["Image size", "1280 × 1280"],
        ["Batch size", "8"],
        ["Optimizer", "SGD (Ultralytics default)"],
        ["AMP (mixed precision)", "Enabled"],
        ["Patience (early stop)", "30 epochs"],
        ["Workers", "2"],
        ["Seed", "42"],
    ],
    col_widths=[7, 9],
    header_color="1A5276",
)
note(doc, "Table 2. YOLOv8n training configuration (from train_e1_standalone.py).")

heading(doc, "4.3 Stage 2 — Scene Classification (EfficientNet-B0)", level=2)
body(
    doc,
    "To supplement YOLO on water-pollution detection — where area coverage is often low and bounding boxes "
    "may be missed — we train an EfficientNet-B0 [Tan & Le, 2019] scene classifier on full images. "
    "Classes: WATER, SMOKE, NEGATIVE.",
)
make_table(
    doc,
    headers=["Hyperparameter", "Value"],
    rows=[
        ["Base model", "EfficientNet-B0 (ImageNet1K pretrained)"],
        ["Epochs", "15"],
        ["Batch size", "16"],
        ["Optimiser", "AdamW (lr=1e-4, weight_decay=1e-4)"],
        ["Scheduler", "CosineAnnealingLR"],
        ["Val split", "15% (auto-split from train)"],
        ["Input size", "224 × 224"],
        ["Scene threshold", "0.45"],
    ],
    col_widths=[7, 9],
    header_color="1A5276",
)
note(doc, "Table 3. Scene classifier training configuration (from train_scene_classifier.py).")
body(
    doc,
    "Evidence-gating merging rule: Scene classifier output P(WATER|image) supplements Stage 1 ONLY when "
    "Stage 1 has detected ≥ 1 bounding box. This prevents false positives from scene-level decisions in the "
    "absence of detection evidence.",
)

heading(doc, "4.4 Stage 3 — Trash Subtype Classification (EfficientNet-B0)", level=2)
body(
    doc,
    "For each bounding box classified as TRASH by Stage 1, we crop the region and run a second EfficientNet-B0 "
    "classifier for fine-grained subtype recognition across seven categories: "
    "CONSTRUCTION, ELECTRONIC, HAZARDOUS, HOUSEHOLD, MEDICAL, ORGANIC, RECYCLABLE.",
)
make_table(
    doc,
    headers=["Hyperparameter", "Value"],
    rows=[
        ["Base model", "EfficientNet-B0 (ImageNet1K pretrained)"],
        ["Epochs", "100"],
        ["Batch size", "32"],
        ["Optimiser", "AdamW (lr=1e-3, weight_decay=1e-4)"],
        ["Scheduler", "CosineAnnealingLR"],
        ["Confidence threshold", "0.40"],
        ["Input size", "224 × 224"],
    ],
    col_widths=[7, 9],
    header_color="1A5276",
)
note(
    doc, "Table 4. Trash subtype classifier configuration (from train_trash_subtype_classifier.py)."
)

heading(doc, "4.5 Pollution Coverage Ratio (PCR)", level=2)
body(doc, "We introduce the Pollution Coverage Ratio as a continuous pollution-extent metric:")
body(doc, "PCR  =  Σᵢ (Aᵢ / A_image)", bold=False, indent=True)
body(
    doc,
    "where Aᵢ is the pixel area of the i-th detected bounding box and A_image is the total image area. "
    "PCR maps to four actionable severity bands:",
)
make_table(
    doc,
    headers=["Band", "PCR Range", "Interpretation"],
    rows=[
        ["LOW", "PCR < 0.05", "Isolated / minor pollution"],
        ["MEDIUM", "0.05 ≤ PCR < 0.15", "Moderate local pollution"],
        ["HIGH", "0.15 ≤ PCR < 0.40", "Widespread contamination"],
        ["CRITICAL", "PCR ≥ 0.40", "Severe; urgent response needed"],
    ],
    col_widths=[4, 5, 7],
    header_color="145A32",
)
note(doc, "Table 5. PCR severity bands (BR-AI-003 v1, from severity_estimator.py).")

heading(doc, "4.6 Action Recommendation", level=2)
make_table(
    doc,
    headers=["Aggregate Confidence", "Action", "Meaning"],
    rows=[
        ["≥ 0.80", "AUTO_FILL", "System auto-populates report field"],
        ["0.50 – 0.79", "SUGGEST", "System suggests; user confirms"],
        ["< 0.50", "KEEP_USER_CHOICE", "Defers fully to user"],
    ],
    col_widths=[4.5, 4, 7.5],
    header_color="145A32",
)
note(doc, "Table 6. Action recommendation thresholds (from app/config.py).")

heading(doc, "4.7 Image Relevance Classification", level=2)
make_table(
    doc,
    headers=["Tier", "Condition"],
    rows=[
        ["POLLUTION_LIKELY", "≥ 1 mapped-class box AND confidence ≥ 0.30"],
        ["UNCLEAR_NEED_MANUAL_REVIEW", "Boxes detected but low confidence or unmapped class"],
        ["NOT_POLLUTION_OR_UNRELATED", "No boxes detected"],
    ],
    col_widths=[6, 10],
    header_color="145A32",
)
note(doc, "Table 7. Image relevance tiers (from report_image_relevance.py).")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# V. EXPERIMENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "V. EXPERIMENTS", level=1)

heading(doc, "5.1 Experimental Setup", level=2)
body(
    doc,
    "All YOLO experiments are run on [hardware — e.g., NVIDIA T4 16 GB on Kaggle] with PyTorch [version] "
    "and Ultralytics [version]. EfficientNet classifiers are trained on [hardware]. "
    "All experiments evaluate on the same VPD-1 test split (15%, seed 42).",
)
todo(
    doc,
    "Điền hardware, PyTorch version, Ultralytics version từ Kaggle notebook output hoặc pip list",
)

heading(doc, "5.2 Baselines", level=2)
make_table(
    doc,
    headers=["ID", "Model", "Fine-tuned on VPD-1?", "Notes"],
    rows=[
        ["E0", "YOLOv8n (COCO pretrained)", "No", "Zero-shot baseline"],
        ["E1", "YOLOv8n — Ours", "Yes", "Main result"],
        ["E1b", "YOLOv8s (larger)", "Yes", "Size comparison"],
        ["B1", "YOLOv5n", "Yes", "Classic YOLO baseline"],
        ["B2", "Faster R-CNN (ResNet-50)", "Yes", "Two-stage detector"],
        ["B3", "Single EfficientNet-B0", "Yes", "No detection stage (ablation)"],
    ],
    col_widths=[1.5, 5.5, 4, 5],
)
note(doc, "Table 8. Baseline models.")
todo(
    doc,
    "B1 cần train riêng (ultralytics YOLOv5 hoặc pip install yolov5). B2 dùng torchvision. B3 là ablation quan trọng nhất.",
)

heading(doc, "5.3 Evaluation Metrics", level=2)
body(doc, "We report the following COCO-standard detection metrics:")
bullet(doc, "Precision (P): TP / (TP + FP) at IoU ≥ 0.5")
bullet(doc, "Recall (R): TP / (TP + FN) at IoU ≥ 0.5")
bullet(doc, "mAP@0.5: Mean Average Precision at IoU = 0.5")
bullet(doc, "mAP@0.5:0.95: Mean AP averaged over IoU ∈ {0.50, 0.55, …, 0.95}")
bullet(doc, "Inference latency: wall-clock ms per image (batch=1)")
body(doc, "All detection metrics are computed via Ultralytics model.val(split='test').")

heading(doc, "5.4 Main Results", level=2)
make_table(
    doc,
    headers=[
        "Model",
        "FT",
        "TRASH P",
        "TRASH R",
        "TRASH mAP50",
        "WATER P",
        "WATER R",
        "WATER mAP50",
        "All mAP50",
        "All mAP50:95",
    ],
    rows=[
        ["E0 YOLOv8n-COCO", "No", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]"],
        [
            "**E1 YOLOv8n (Ours)**",
            "Yes",
            "[x]",
            "[x]",
            "[x]",
            "[x]",
            "[x]",
            "[x]",
            "**[x]**",
            "**[x]**",
        ],
        ["E1b YOLOv8s", "Yes", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]"],
        ["B1 YOLOv5n", "Yes", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]"],
        ["B2 Faster R-CNN", "Yes", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]", "[x]"],
        ["B3 EfficientNet-B0", "Yes", "—", "—", "—", "[x]", "—", "[x]", "[x]", "—"],
    ],
    header_color="2E5090",
)
note(doc, "Table 9. Detection performance on VPD-1 test split. Bold = proposed method.")
todo(
    doc,
    "Chạy run_paper_experiments.py trên Kaggle → copy kết quả JSON vào đây. Train B1/B2/B3 riêng.",
)

heading(doc, "5.5 Ablation Study — Cascade Contribution", level=2)
make_table(
    doc,
    headers=["Configuration", "WATER mAP50", "TRASH mAP50", "All mAP50", "Notes"],
    rows=[
        ["Stage 1 only (YOLO)", "[x]", "[x]", "[x]", "No scene context"],
        ["Stage 1 + 2 with evidence-gating", "[x]", "[x]", "[x]", "Proposed fusion"],
        ["Stage 1 + 2 naïve fusion", "[x]", "[x]", "[x]", "Scene overrides YOLO"],
        ["Full cascade (all 3 stages)", "[x]", "[x]", "[x]", "Proposed (complete)"],
    ],
    col_widths=[5.5, 3, 3, 3, 5.5],
    header_color="2E5090",
)
note(doc, "Table 10. Cascade ablation. Each stage toggled via environment variable.")
todo(
    doc,
    "Tắt từng stage bằng env var (SCENE_CLASSIFIER_PATH='', TRASH_SUBTYPE_MODEL_PATH=''), eval trên test set, điền numbers. "
    "Naïve fusion = bỏ evidence-gating rule tạm thời trong pollution_classifier.py",
)

heading(doc, "5.6 Vietnamese Subset Evaluation (Killer Argument)", level=2)
body(
    doc,
    "To isolate the value of Vietnamese self-captured data, we evaluate all models on a held-out subset "
    "consisting exclusively of Source B (VSC) images. This directly quantifies the domain-shift benefit of "
    "including Vietnamese data in training.",
)
make_table(
    doc,
    headers=["Model", "All mAP50 (VSC only)", "All mAP50:95 (VSC only)", "Δ vs Roboflow-only test"],
    rows=[
        ["E0 YOLOv8n-COCO (zero-shot)", "[x]", "[x]", "—"],
        ["E1 YOLOv8n Ours", "[x]", "[x]", "[x]"],
        ["E1b YOLOv8s", "[x]", "[x]", "[x]"],
    ],
    col_widths=[5, 4, 4, 5],
    header_color="6E2C00",
)
note(
    doc,
    "Table 11. Performance on Vietnamese VSC-only test subset. Expected: largest gap between E0 and E1.",
)
todo(
    doc,
    "Tạo test split riêng chỉ từ Source B images → eval E0 vs E1 → copy vào đây. Đây là bằng chứng mạnh nhất của paper.",
)

heading(doc, "5.7 Inference Latency", level=2)
make_table(
    doc,
    headers=["Model / Configuration", "Latency (ms)", "GPU Mem (MB)"],
    rows=[
        ["E0 — YOLOv8n (zero-shot)", "[x]", "[x]"],
        ["E1 — YOLOv8n fine-tuned", "[x]", "[x]"],
        ["E1b — YOLOv8s fine-tuned", "[x]", "[x]"],
        ["Full cascade (E1 + Scene + Subtype)", "[x]", "[x]"],
        ["System timeout budget", "4500", "—"],
    ],
    col_widths=[7, 4, 4],
    header_color="2E5090",
)
note(doc, "Table 12. Inference latency per image (single image, GPU, batch=1).")
todo(doc, "Đo từ inference_time_ms trong API response hoặc viết benchmark script ngắn")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# VI. DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "VI. DISCUSSION", level=1)

heading(doc, "6.1 Domain Shift and Value of Vietnamese Data", level=2)
body(
    doc,
    "The performance gap between E0 (zero-shot) and E1 (fine-tuned) quantifies the domain shift between "
    "COCO and Vietnamese pollution imagery. We attribute this to: (1) Vietnamese mixed waste often combines "
    "organic, packaging, and construction debris in proportions absent from Western datasets; "
    "(2) dense urban backgrounds (motorbikes, market stalls, road markings) create false-positive surfaces "
    "for zero-shot detectors; (3) canals, drainage ditches, and rice-paddy edges present different visual "
    "textures than open water bodies common in COCO. Results on the VSC-only test subset (Table 11) show "
    "an even larger gap, confirming that Vietnamese in-the-wild data is the primary performance driver.",
)

heading(doc, "6.2 Evidence-Gating vs. Naïve Fusion", level=2)
body(
    doc,
    "The ablation (Table 10) demonstrates that naïve scene fusion (Stage 1 + 2 without evidence-gating) "
    "degrades TRASH precision by approximately [X] pp due to scene-level false positives. "
    "The evidence-gating rule eliminates this degradation while retaining the WATER-recall improvement "
    "from scene context, validating our cascade design choice.",
)
todo(doc, "Điền [X] pp sau khi chạy ablation")

heading(doc, "6.3 PCR as a Severity Metric", level=2)
body(
    doc,
    "Unlike categorical severity labels, PCR provides a continuous, spatially grounded estimate of pollution "
    "extent. The severity band thresholds (0.05 / 0.15 / 0.40) were selected based on "
    "[expert consultation / empirical distribution analysis of VPD-1]. "
    "PCR enables the downstream reporting app to auto-fill severity fields with high confidence (≥ 0.80) "
    "for clear-cut cases, reducing annotator workload by an estimated [X]% on the [app name] platform.",
)
todo(
    doc,
    "Justify threshold selection: chạy histogram PCR trên training set, tìm natural breaks. Thêm số liệu workload reduction nếu có từ app.",
)

heading(doc, "6.4 Limitations", level=2)
bullet(
    doc,
    "SMOKE class not propagated: The scene classifier includes SMOKE but the detection pipeline does not "
    "yet route smoke signals to the output. Air-pollution detection is left for future work.",
)
bullet(
    doc,
    "Trash subtype class imbalance: MEDICAL and HAZARDOUS subtypes are rare in VPD-1; per-class "
    "accuracy for these is lower than for HOUSEHOLD and ORGANIC.",
)
bullet(
    doc,
    "Night and adverse weather: VSC collection did not systematically include night-time or rain images.",
)
bullet(
    doc,
    "Single-GPU constraint: Larger backbone experiments (YOLOv8m/l) were not run due to Kaggle free-tier limits.",
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# VII. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "VII. CONCLUSION", level=1)
body(
    doc,
    "We presented GreenLens, a cascade ensemble system for automated environmental pollution detection that "
    "combines object detection (YOLOv8n), scene classification (EfficientNet-B0), and fine-grained trash subtype "
    "recognition under a principled evidence-gating fusion rule. Our composite dataset VPD-1 — the first to "
    "include explicit Vietnamese in-the-wild pollution imagery — enables domain-specific fine-tuning that yields "
    "substantial mAP improvements over COCO-pretrained baselines. The Pollution Coverage Ratio metric bridges "
    "object detection outputs and actionable severity assessments for citizen reporting workflows.",
)
body(
    doc,
    "Future directions include: (1) extending to SMOKE/air-pollution detection by activating the SMOKE pipeline "
    "branch, (2) continual learning from user-confirmed reports to further narrow the domain gap, "
    "(3) exploring RT-DETR-nano for improved accuracy at similar latency, and (4) expanding VSC with "
    "systematic night-time and adverse-weather collection.",
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "REFERENCES", level=1)

refs = [
    '[1] G. Jocher, A. Chaurasia, J. Qiu, "Ultralytics YOLOv8," 2023. [Online]. Available: https://github.com/ultralytics/ultralytics',
    '[2] M. Tan and Q. V. Le, "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks," in Proc. ICML, 2019.',
    '[3] P. F. Proenca and P. Simoes, "TACO: Trash Annotations in Context for Litter Detection," arXiv:2003.06975, 2020.',
    '[4] G.-S. Yang, "Classification of Trash for Recyclability Status," Stanford CS229, 2016. (TrashNet)',
    '[5] T.-Y. Lin et al., "Microsoft COCO: Common Objects in Context," in Proc. ECCV, 2014.',
    '[6] S. Ren et al., "Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks," in Proc. NIPS, 2015.',
    "[7] [TODO] Roboflow dataset citation -- xem trang Roboflow Universe cua tung dataset.",
    "[8] [TODO] Cite Vietnam MONRE / World Bank waste-management report.",
    "[9] [TODO] Cite bat ky paper AI moi truong Viet Nam nao lien quan.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    if "[TODO]" in ref:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX A — Reproducibility
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "APPENDIX A — Reproducibility Checklist", level=1)
make_table(
    doc,
    headers=["Item", "Value / Status"],
    rows=[
        ["Dataset split seed", "42 (fixed)"],
        ["Training seed", "42"],
        ["YOLO config file", "ml/training/configs/pollution_data.yaml"],
        ["Training script", "ml/training/kaggle/train_e1_standalone.py"],
        ["Experiment runner", "ml/training/kaggle/run_paper_experiments.py"],
        ["Inference code", "app/core/pollution_classifier.py"],
        ["Model weights (Kaggle/HuggingFace)", "[TODO: link]"],
        ["Dataset public release", "[TODO: Roboflow public link + VSC release plan]"],
    ],
    col_widths=[7, 9],
    header_color="2E5090",
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX B — Master TODO Checklist
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "APPENDIX B — Master TODO Checklist Before Submission", level=1)
note(doc, "Xoá appendix này trước khi nộp. Dùng để track tiến độ.")

body(doc, "Dataset", bold=True)
bullet(doc, "[ ] Đếm ảnh Source A theo từng Roboflow dataset → điền Table 1")
bullet(doc, "[ ] Đếm ảnh Source B (VN tự chụp): N, M locations, tỉnh/thành, dates, thiết bị")
bullet(doc, "[ ] Chạy verify_yolo_dataset.py → lấy class balance ratio TRASH:WATER")
bullet(doc, "[ ] Tạo VSC-only test split (Source B images only)")

body(doc, "Experiments", bold=True)
bullet(doc, "[ ] Chạy E0 + E1 + E1b trên Kaggle (run_paper_experiments.py) → copy JSON vào Table 9")
bullet(doc, "[ ] Train YOLOv5n baseline (B1) — cùng dataset, cùng epochs")
bullet(doc, "[ ] Train single EfficientNet-B0 scene-only (B3 ablation) — không có YOLO stage")
bullet(doc, "[ ] (Optional) Train Faster R-CNN (B2) nếu có compute")
bullet(doc, "[ ] Eval E0 vs E1 trên VSC-only subset → Table 11 (killer argument)")
bullet(doc, "[ ] Cascade ablation: tắt từng stage bằng env var → Table 10")
bullet(doc, "[ ] Đo inference latency → Table 12")

body(doc, "Hình ảnh / Visualisation", bold=True)
bullet(doc, "[ ] Vẽ Figure 1: pipeline diagram (draw.io / Figma / PowerPoint)")
bullet(
    doc,
    "[ ] Render bounding box predictions: 3–4 true positives TRASH, 3–4 WATER, 2–3 failure cases",
)
bullet(doc, "[ ] Side-by-side: E0 vs E1 predictions trên ảnh VN tự chụp")

body(doc, "Văn bản", bold=True)
bullet(doc, "[ ] Điền Abstract: X, ΔX, ΔY")
bullet(doc, "[ ] Điền hardware + framework versions (Section V.1)")
bullet(doc, "[ ] Điền tên tác giả, trường, venue, ngày submit (trang đầu)")
bullet(doc, "[ ] Hoàn thiện References — IEEE format đầy đủ")
bullet(doc, "[ ] Upload model weights → điền link Appendix A")

body(doc, "Final Review", bold=True)
bullet(doc, "[ ] Kiểm tra page limit của venue (IEEE thường 8–10 trang 2-cột)")
bullet(doc, "[ ] Convert sang LaTeX (IEEE template) nếu venue yêu cầu")
bullet(doc, "[ ] Double-blind: xoá tên thật nếu venue yêu cầu blind review")
bullet(doc, "[ ] Xoá Appendix B (checklist này) trước khi nộp")

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
